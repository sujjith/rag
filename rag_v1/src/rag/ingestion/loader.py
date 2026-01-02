
from pathlib import Path  # Modern way to handle file paths
import fitz  # PyMuPDF - the package name is 'pymupdf' but import is 'fitz'
from docx import Document as DocxDocument  # Renamed to avoid conflict with our Document
from rag.models import Document  # Our Document dataclass from models.py (not python-docx)
import hashlib  # For generating unique document IDs from file paths

class DocumentLoader:  # Load documents from various file formats (PDF, TXT, MD, DOCX)

    def load(self, file_path: Path) -> Document:  # Main method - load a document based on file extension
        suffix = file_path.suffix.lower()  # Get file extension in lowercase (e.g., ".pdf", ".txt")

        loaders = {  # Map file extensions to their loader methods (avoids long if/elif chains)
            ".pdf": self._load_pdf,
            ".txt": self._load_text,
            ".md": self._load_text,  # Markdown uses same loader as plain text
            ".docx": self._load_docx,
        }

        loader = loaders.get(suffix)  # Get loader function for this file type (returns None if not found)
        if not loader:
            raise ValueError(f"Unsupported file type: {suffix}")

        content = loader(file_path)  # Extract text content using the appropriate loader
        
        doc_id = hashlib.md5(str(file_path).encode()).hexdigest()[:12]  # Generate unique ID using MD5 hash (same file = same ID)

        return Document(  # Create and return a Document object with extracted info
            id=doc_id,
            content=content,
            source=str(file_path),
            metadata={"filename": file_path.name, "type": suffix}
        )

    def _load_pdf(self, path: Path) -> str:  # Extract text from PDF
        with fitz.open(path) as doc:  # 'with' ensures PDF is properly closed after reading
            return "\n".join(page.get_text() for page in doc)  # Extract text from each page, join with newlines

    def _load_text(self, path: Path) -> str:  # Load plain text file
        return path.read_text(encoding="utf-8")  # Simply read the entire file as UTF-8 text

    def _load_docx(self, path: Path) -> str:  # Extract text from DOCX
        doc = DocxDocument(path)  # Open Word document
        return "\n".join(para.text for para in doc.paragraphs)  # Extract text from all paragraphs (not tables/headers/footers)
